#!/usr/bin/env python3
"""
build_docx.py <input.md> <output.docx> [--cover cover.png]

Modern-editorial styled book builder. Post-pass does:
  - run-level heading color (forces ONE navy; defeats Google Docs theme override)
  - section numbering 1, 1.1, 1.1.1 on headings
  - chapter abstract (first italic para after a Chapter H1) -> light-blue box
  - callout blockquotes (WORKED EXAMPLE / BEYOND VERMONT / TRY THIS) -> shaded boxes
  - "Sources" body text -> small font
  - tables: full width portrait; WIDE tables (>=5 col) get SMALLER font, NO landscape
  - extra space after every table
  - cover image on page 1
  - TOC on its OWN page; page break before each chapter
  - strip Pandoc heading bookmark anchors (the blue ribbon artifacts)
"""
import sys, subprocess, os, re, copy

BODY_FONT = "Garamond"   # must match SERIF in make_reference.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor

GRAY = RGBColor(0x40, 0x40, 0x40)   # Sources / citation grey
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

args = sys.argv[1:]
SRC, OUT = args[0], args[1]
COVER = args[args.index('--cover')+1] if '--cover' in args else None
HERE = os.path.dirname(os.path.abspath(__file__))
REF = os.path.join(HERE,'reference.docx')

NAVY = RGBColor(0x1B,0x3A,0x6B)
HEADCOLOR = RGBColor(0x16,0x1B,0x22)   # near-black book headings (match make_reference.py)
PORTRAIT_DXA = int(6.3*1440)   # full text width (portrait, 8.5in - 2*1.1in margins)
LANDSCAPE_DXA = int(9.6*1440)  # full text width on a landscape page (11in - 2*0.7in margins)
WIDE_TABLE_MIN_COLS = 4        # a table must have >= this many columns ...
WIDE_TABLE_MIN_CELL = 120      # ... AND at least one cell this long (chars) to go landscape

def _is_dependency_matrix(t):
    """The Figure 1.B six-pillar dependency matrix, identified by its unmistakable
    top-left header cell. It is 7 columns of short cells, so it does NOT trip the
    generic dense-cell test, but it is far too wide for portrait — force it
    landscape explicitly."""
    try:
        return t.rows[0].cells[0].text.strip().startswith('From')
    except Exception:
        return False

def _table_is_wide_dense(t):
    """A table qualifies for landscape only if it is BOTH wide (>= min cols) AND
    dense (has at least one genuinely long cell). Sparse short-celled 4-col tables
    read fine in portrait and should stay upright. The Figure 1.B dependency
    matrix is an explicit exception — wide but short-celled — forced landscape."""
    if len(t.rows) == 1 and len(t.columns) == 1:
        return False
    if _is_dependency_matrix(t):
        return False   # PORTRAIT — author finishes matrix formatting manually
    if len(t.columns) < WIDE_TABLE_MIN_COLS:
        return False
    maxlen = max((len(c.text) for r in t.rows for c in r.cells), default=0)
    return maxlen >= WIDE_TABLE_MIN_CELL

# Callout palette: (fill bg, left-bar color, label color)
CALLOUTS = {
    'WORKED EXAMPLE': ('FFF6E5','E0991A','8A5A00'),   # amber
    'BEYOND VERMONT': ('E7F4F2','1F8A78','0F5247'),   # teal
    'TRY THIS':       ('ECEEFB','4759C9','27317A'),   # indigo
    'KEY POINT':      ('F0F1F4','5B6B86','2C3E55'),   # slate
    'VERMONT IN PRACTICE': ('EAF3EA','3E7B3E','255025'),  # Vermont green
}

subprocess.run(['pandoc',SRC,'-o',OUT,'--reference-doc',REF,
    '--resource-path', f'{HERE}:.',
    '--from','markdown+pipe_tables+grid_tables+raw_html-raw_attribute','--wrap','none'],check=True)
doc = Document(OUT)
body = doc.element.body

def set_run_color(p, rgb):
    for r in p.runs:
        r.font.color.rgb = rgb

def shade(el, hexfill):
    """Apply shading to a tc or p element via pPr/tcPr."""
    pr = el
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexfill)
    pr.append(shd)

def strip_heading_bookmarks():
    """Remove Pandoc-injected bookmarkStart/End (blue ribbons in GDocs), then add
    clean _Toc-named bookmarks around every heading so the live TOC field can
    resolve real page numbers. _Toc-prefixed bookmarks do NOT render as visible
    ribbons in Google Docs, but the TOC field (\\h \\u) uses them as anchors."""
    for bm in list(body.iter(qn('w:bookmarkStart'))):
        bm.getparent().remove(bm)
    for bm in list(body.iter(qn('w:bookmarkEnd'))):
        bm.getparent().remove(bm)
    # now wrap each heading in a fresh _Toc bookmark
    _bid = [1000]
    for p in doc.paragraphs:
        if p.style.name not in ('Heading 1','Heading 2','Heading 3'):
            continue
        if not p.text.strip():
            continue
        name = f'_Toc{_bid[0]}'; _bid[0]+=1
        p._p.set('data-toc-bookmark', name) if False else None
        start = OxmlElement('w:bookmarkStart'); start.set(qn('w:id'), str(_bid[0]))
        start.set(qn('w:name'), name)
        end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), str(_bid[0]))
        # insert start after pPr (before first run) and end at paragraph tail
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None: pPr.addnext(start)
        else: p._p.insert(0, start)
        p._p.append(end)
        _HEADING_BOOKMARKS.append((p.style.name, p.text.strip(), name))

_HEADING_BOOKMARKS = []  # (style, text, bookmark_name) filled by strip_heading_bookmarks()

# ── 1. headings: force navy at run level + section numbering ─────────────────
def _first_run_index(p_el):
    idx=0
    for i,ch in enumerate(p_el):
        if ch.tag==qn('w:pPr'): idx=i+1
    return idx

def _prefix(p, text):
    run_el = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    col = OxmlElement('w:color'); col.set(qn('w:val'),'161B22'); rpr.append(col)  # near-black heading number
    run_el.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text=text; run_el.append(t)
    p._p.insert(_first_run_index(p._p), run_el)

def number_and_color_headings():
    """Color all headings navy. Number H2/H3 ONLY inside numbered Chapters/Appendices,
    chapter-relative (1.1, 1.2.1). Front matter (Preface/Intro/Conclusion/TOC) unnumbered."""
    chap=None; c=[0,0]
    for p in doc.paragraphs:
        st=p.style.name
        if st=='Heading 1':
            set_run_color(p, HEADCOLOR)
            m=re.match(r'\s*(Chapter|Appendix)\s+([0-9A-Z]+)', p.text)
            chap = m.group(2) if m else None
            c=[0,0]
        elif st=='Heading 2':
            set_run_color(p, HEADCOLOR)
            if chap:
                c[0]+=1; c[1]=0
                _prefix(p, f"{chap}.{c[0]}  ")
        elif st=='Heading 3':
            set_run_color(p, HEADCOLOR)
            if chap:
                c[1]+=1
                _prefix(p, f"{chap}.{c[0]}.{c[1]}  ")

# ── 2. chapter abstract -> light blue box ────────────────────────────────────
def style_chapter_abstracts():
    paras = doc.paragraphs
    for i,p in enumerate(paras):
        if p.style.name=='Heading 1' and re.match(r'\s*Chapter', p.text):
            # next non-empty paragraph that is fully italic = abstract
            j=i+1
            while j<len(paras) and not paras[j].text.strip(): j+=1
            if j<len(paras):
                ab=paras[j]
                is_ital = ab.runs and all((r.italic or not r.text.strip()) for r in ab.runs)
                if is_ital:
                    box_paragraph(ab, fill='E8F0FB', bar='2457A4', textcolor='1B3A6B', italic=True)

# ── 3. callouts: fenced-div Callout* styled paragraph groups -> shaded box ────
# Each callout is a run of consecutive paragraphs that all share a Callout* style.
# We wrap the whole run in ONE single-cell shaded table so the background covers
# the entire callout (label + body). Body font is smaller + italic to differ.
CALLOUT_STYLES = {
    'CalloutWorked': CALLOUTS['WORKED EXAMPLE'],
    'CalloutBeyond': CALLOUTS['BEYOND VERMONT'],
    'CalloutTry':    CALLOUTS['TRY THIS'],
    'CalloutKey':    CALLOUTS['KEY POINT'],
    'CalloutVT':     CALLOUTS['VERMONT IN PRACTICE'],
}

def _is_bullet(ch):
    """A Pandoc list item inside a div: Normal style with a numPr."""
    if ch.tag != qn('w:p'): return False
    return ch.find('.//'+qn('w:numPr')) is not None

def _is_empty_p(ch):
    if ch.tag != qn('w:p'): return False
    return not ''.join(t.text or '' for t in ch.iter(qn('w:t'))).strip()

def style_key_concepts():
    """Within every 'Key Concepts' section, style each glossary entry to match the
    author's Example #7: TERM = Calibri bold blue (2E74B5) 10pt; DEFINITION =
    Calibri regular black 9pt. A term is a bold-only line; the definition is the
    paragraph(s) under it, until the next term or the next heading."""
    paras=doc.paragraphs
    n=len(paras); i=0
    def is_kc_heading(p): return p.style.name.startswith('Heading') and 'Key Concept' in p.text
    def is_heading(p): return p.style.name.startswith('Heading')
    while i<n:
        if is_kc_heading(paras[i]):
            j=i+1
            while j<n and not is_heading(paras[j]):
                p=paras[j]
                if not p.text.strip(): j+=1; continue
                # Each entry is one paragraph: a leading BOLD run (the term) +
                # a hard line break + non-bold run(s) (the definition).
                # Style per-run: bold run -> blue 10pt term; rest -> black 9pt def.
                for r in p.runs:
                    r.font.name='Calibri'
                    if r.bold:
                        r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x2E,0x74,0xB5)
                    else:
                        r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x00,0x00,0x00)
                pf=p.paragraph_format; pf.space_before=Pt(8); pf.space_after=Pt(6)
                j+=1
            i=j
        else:
            i+=1

def style_depmap():
    """Render the DepMap block (the dependency-map replacement for the old 6x6
    matrix). Groups of consecutive DepMap paragraphs become one light-tinted box.
    A pillar header line (starts with a bold ALLCAPS pillar name, e.g. 'POLICY —')
    is styled navy bold; arrow lines ('→ Target · [TYPE] ...') get their [TYPE]
    tag colored by relationship kind and the → arrow in navy."""
    TYPE_COLOR = {
        'ENABLES':    'C97A0F',  # amber
        'REQUIRES':   '1F6B5D',  # teal
        'DRIVES':     '3E4FB0',  # indigo
        'CONSTRAINS': '9A2A48',  # maroon
    }
    children=list(body.iterchildren())
    i=0
    while i<len(children):
        if _para_style(children[i])=='DepMap':
            grp=[children[i]]; j=i+1
            while j<len(children) and (_para_style(children[j])=='DepMap' or _is_empty_p(children[j])):
                if _para_style(children[j])=='DepMap': grp.append(children[j])
                j+=1
            _build_depmap_box(grp, TYPE_COLOR)
            i=i+(j-i)
        else:
            i+=1

def _build_depmap_box(paras, TYPE_COLOR):
    import re as _re
    first=paras[0]
    # one single-cell shaded table so the tint spans the whole block
    tbl=OxmlElement('w:tbl'); tblPr=OxmlElement('w:tblPr')
    tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(PORTRAIT_DXA)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    bd=OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right'):
        b=OxmlElement(f'w:{e}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'C9D2DD'); bd.append(b)
    tblPr.append(bd)
    cm=OxmlElement('w:tblCellMar')
    for e,w in (('top','120'),('left','160'),('bottom','120'),('right','160')):
        m=OxmlElement(f'w:{e}'); m.set(qn('w:w'),w); m.set(qn('w:type'),'dxa'); cm.append(m)
    tblPr.append(cm); tbl.append(tblPr)
    grid=OxmlElement('w:tblGrid'); gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(PORTRAIT_DXA)); grid.append(gc); tbl.append(grid)
    tr=OxmlElement('w:tr'); tc=OxmlElement('w:tc')
    tcPr=OxmlElement('w:tcPr'); tcW=OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(PORTRAIT_DXA)); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F5F7FA'); tcPr.append(shd); tc.append(tcPr)
    first.addprevious(tbl)
    for pe in paras:
        pe.getparent().remove(pe)
        pPr=pe.find(qn('w:pPr'))
        if pPr is None: pPr=OxmlElement('w:pPr'); pe.insert(0,pPr)
        for old in pPr.findall(qn('w:pStyle')): pPr.remove(old)
        for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
        txt=''.join(t.text or '' for t in pe.iter(qn('w:t')))
        is_header = bool(_re.match(r'^[A-Z][A-Z]+\s+—', txt.strip()))
        sp=OxmlElement('w:spacing')
        sp.set(qn('w:before'), '120' if is_header else '10'); sp.set(qn('w:after'),'10')
        sp.set(qn('w:line'),'240'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
        # restyle every run: base Calibri 9.5; header navy bold; color the [TYPE] tag
        for r in pe.findall(qn('w:r')):
            rpr=r.find(qn('w:rPr'))
            if rpr is None: rpr=OxmlElement('w:rPr'); r.insert(0,rpr)
            rf=rpr.find(qn('w:rFonts'))
            if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
            rf.set(qn('w:ascii'),'Calibri'); rf.set(qn('w:hAnsi'),'Calibri')
            for tag in ('w:sz','w:szCs'):
                e=rpr.find(qn(tag))
                if e is None: e=OxmlElement(tag); rpr.append(e)
                e.set(qn('w:val'),'19')  # 9.5pt
            t=r.find(qn('w:t')); rtext=t.text if t is not None else ''
            if is_header:
                col=rpr.find(qn('w:color'))
                if col is None: col=OxmlElement('w:color'); rpr.append(col)
                col.set(qn('w:val'),'1B3A6B')
            else:
                # color a [TYPE] tag run
                m=_re.search(r'\[(ENABLES|REQUIRES|DRIVES|CONSTRAINS)\]', rtext or '')
                if m:
                    col=rpr.find(qn('w:color'))
                    if col is None: col=OxmlElement('w:color'); rpr.append(col)
                    col.set(qn('w:val'),TYPE_COLOR[m.group(1)])
                    if rpr.find(qn('w:b')) is None: rpr.append(OxmlElement('w:b'))
        tc.append(pe)
    tr.append(tc); tbl.append(tr)
    sp_after=OxmlElement('w:p'); pr=OxmlElement('w:pPr'); s=OxmlElement('w:spacing'); s.set(qn('w:after'),'120'); pr.append(s); sp_after.append(pr)
    tbl.addnext(sp_after)

def style_banners():
    """Render Banner paragraphs as a full-width light-tint bar: pale navy fill,
    navy bold text, thick navy top+bottom accent rules, centered. A blank spacer
    paragraph is inserted before AND after so the tinted bar never touches the
    surrounding body text."""
    for p in doc.paragraphs:
        if p.style.name!='Banner': continue
        pPr=p._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:shd')): pPr.remove(old)
        shade(pPr,'E8F0FB')   # pale navy tint
        pbdr=OxmlElement('w:pBdr')
        for edge,(sz,col) in (('top',('8','1B3A6B')),('bottom',('8','1B3A6B')),
                              ('left',('4','E8F0FB')),('right',('4','E8F0FB'))):   # thinner rules (was 18)
            b=OxmlElement(f'w:{edge}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz)
            b.set(qn('w:space'),'6'); b.set(qn('w:color'),col); pbdr.append(b)
        pPr.append(pbdr)
        jc=OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
        sp=OxmlElement('w:spacing'); sp.set(qn('w:before'),'40'); sp.set(qn('w:after'),'40')
        sp.set(qn('w:line'),'240'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
        for r in p.runs:
            r.font.name='Calibri'; r.font.size=Pt(11); r.bold=True   # smaller (was 15)
            r.font.color.rgb=RGBColor(0x1B,0x3A,0x6B); r.italic=False
        # insert un-shaded blank spacer paragraphs before and after the bar so the
        # tint/border never abuts the body text.
        def _blank_spacer():
            spp=OxmlElement('w:p'); sppr=OxmlElement('w:pPr')
            s=OxmlElement('w:spacing'); s.set(qn('w:before'),'0'); s.set(qn('w:after'),'0')
            s.set(qn('w:line'),'160'); s.set(qn('w:lineRule'),'auto'); sppr.append(s)
            spp.append(sppr); return spp
        p._p.addprevious(_blank_spacer())
        p._p.addnext(_blank_spacer())

def style_part_divider():
    """Render PartDivider paragraphs as a strong full-width navy bar: navy fill,
    white bold uppercase text, thick navy rules top+bottom, centered. Signals a
    major structural break within a chapter (e.g. Part One -> Part Two)."""
    for p in doc.paragraphs:
        if p.style.name!='PartDivider': continue
        pPr=p._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:shd')): pPr.remove(old)
        shade(pPr,'1B3A6B')   # solid navy fill
        pbdr=OxmlElement('w:pBdr')
        for edge in ('top','bottom'):
            b=OxmlElement(f'w:{edge}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'18')
            b.set(qn('w:space'),'8'); b.set(qn('w:color'),'1B3A6B'); pbdr.append(b)
        pPr.append(pbdr)
        jc=OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pPr.append(jc)
        sp=OxmlElement('w:spacing'); sp.set(qn('w:before'),'240'); sp.set(qn('w:after'),'160')
        sp.set(qn('w:line'),'300'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
        for r in p.runs:
            r.font.name='Calibri'; r.font.size=Pt(13); r.bold=True
            r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.italic=False
            # letter-spacing for the "part heading" feel
            rPr=r._r.get_or_add_rPr()
            sp2=OxmlElement('w:spacing'); sp2.set(qn('w:val'),'40'); rPr.append(sp2)
        # a part divider should START A NEW PAGE (not float mid-page in whitespace),
        # and stay glued to the text that follows it.
        pb=OxmlElement('w:pageBreakBefore'); pPr.append(pb)
        kn=OxmlElement('w:keepNext'); pPr.append(kn)
        # remove any blank spacer paragraph immediately before it (it caused the
        # "floating in the middle of nowhere" gap at the bottom of the prior page).

def style_quotes():
    """Force run-level formatting on PullQuote/QuoteAttr paragraphs so the look
    actually applies (style inheritance from Pandoc divs is unreliable).
    PullQuote: italic Garamond 11pt. QuoteAttr: roman Garamond 9.5pt gray."""
    for p in doc.paragraphs:
        st=p.style.name
        if st=='PullQuote':
            for r in p.runs:
                r.italic=True; r.font.name='Garamond'; r.font.size=Pt(11)
                r.font.color.rgb=RGBColor(0x1A,0x1A,0x1A)
        elif st=='QuoteAttr':
            for r in p.runs:
                r.italic=False; r.bold=False; r.font.name='Garamond'; r.font.size=Pt(9.5)
                r.font.color.rgb=RGBColor(0x33,0x33,0x33)

def style_stat_strips():
    """Render each run of consecutive StatStrip paragraphs ('**VALUE** — Label')
    as ONE horizontal card table: N equal columns, big navy value over gray label."""
    children=list(body.iterchildren())
    i=0
    while i<len(children):
        if _para_style(children[i])=='StatStrip':
            grp=[children[i]]; j=i+1
            while j<len(children) and (_para_style(children[j])=='StatStrip' or _is_empty_p(children[j])):
                if _para_style(children[j])=='StatStrip': grp.append(children[j])
                j+=1
            _build_stat_cards(grp); i=i+ (j-i)
        else:
            i+=1

def _build_stat_cards(paras):
    cards=[]
    for pe in paras:
        txt=''.join(t.text or '' for t in pe.iter(qn('w:t')))
        if '—' in txt: val,lbl=txt.split('—',1)
        elif ' - ' in txt: val,lbl=txt.split(' - ',1)
        else: val,lbl=txt,''
        cards.append((val.strip().strip('*'), lbl.strip()))
    n=len(cards) or 1
    first=paras[0]
    tbl=OxmlElement('w:tbl')
    tblPr=OxmlElement('w:tblPr')
    tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(PORTRAIT_DXA)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    # marker so style_tables() skips this (already fully styled as a stat strip)
    cap=OxmlElement('w:tblCaption'); cap.set(qn('w:val'),'STATSTRIP'); tblPr.append(cap)
    bd=OxmlElement('w:tblBorders')
    # white gaps between the light-blue cards (like separate chips); no outer frame
    for e in ('insideV',):
        b=OxmlElement(f'w:{e}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'FFFFFF'); bd.append(b)
    tblPr.append(bd)
    cm=OxmlElement('w:tblCellMar')
    for e,w in (('top','50'),('left','80'),('bottom','50'),('right','80')):   # tighter top/bottom
        m=OxmlElement(f'w:{e}'); m.set(qn('w:w'),w); m.set(qn('w:type'),'dxa'); cm.append(m)
    tblPr.append(cm); tbl.append(tblPr)
    grid=OxmlElement('w:tblGrid'); each=int(PORTRAIT_DXA/n)
    for _ in range(n):
        gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(each)); grid.append(gc)
    tbl.append(grid)
    tr=OxmlElement('w:tr')
    for val,lbl in cards:
        tc=OxmlElement('w:tc')
        tcPr=OxmlElement('w:tcPr'); tcW=OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(each)); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)
        # light-blue cell fill to match the chapter-abstract box (E8F0FB)
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'E8F0FB'); tcPr.append(shd)
        val_align=OxmlElement('w:vAlign'); val_align.set(qn('w:val'),'center'); tcPr.append(val_align)
        tc.append(tcPr)
        # value paragraph: navy bold centered, 8pt, NOT italic
        pv=OxmlElement('w:p'); pvpr=OxmlElement('w:pPr')
        ja=OxmlElement('w:jc'); ja.set(qn('w:val'),'center'); pvpr.append(ja)
        sp=OxmlElement('w:spacing'); sp.set(qn('w:after'),'20'); sp.set(qn('w:before'),'40'); pvpr.append(sp); pv.append(pvpr)
        rv=OxmlElement('w:r'); rvpr=OxmlElement('w:rPr')
        rf=OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Calibri'); rf.set(qn('w:hAnsi'),'Calibri'); rvpr.append(rf)
        bb=OxmlElement('w:b'); rvpr.append(bb)
        szv=OxmlElement('w:sz'); szv.set(qn('w:val'),'16'); rvpr.append(szv)   # 8pt bold
        cv=OxmlElement('w:color'); cv.set(qn('w:val'),'1B3A6B'); rvpr.append(cv)  # dark navy on light blue
        rv.append(rvpr); tv=OxmlElement('w:t'); tv.text=val; rv.append(tv); pv.append(rv); tc.append(pv)
        # label paragraph: navy bold centered, 8pt, NOT italic
        pl=OxmlElement('w:p'); plpr=OxmlElement('w:pPr')
        ja2=OxmlElement('w:jc'); ja2.set(qn('w:val'),'center'); plpr.append(ja2)
        sp2=OxmlElement('w:spacing'); sp2.set(qn('w:after'),'40'); plpr.append(sp2); pl.append(plpr)
        rl=OxmlElement('w:r'); rlpr=OxmlElement('w:rPr')
        rf2=OxmlElement('w:rFonts'); rf2.set(qn('w:ascii'),'Calibri'); rf2.set(qn('w:hAnsi'),'Calibri'); rlpr.append(rf2)
        blb=OxmlElement('w:b'); rlpr.append(blb)                                # bold
        szl=OxmlElement('w:sz'); szl.set(qn('w:val'),'16'); rlpr.append(szl)   # 8pt
        cl=OxmlElement('w:color'); cl.set(qn('w:val'),'1B3A6B'); rlpr.append(cl)  # dark navy on light blue
        rl.append(rlpr); tl=OxmlElement('w:t'); tl.text=lbl; rl.append(tl); pl.append(rl); tc.append(pl)
        tr.append(tc)
    tbl.append(tr)
    first.addprevious(tbl)
    for pe in paras: pe.getparent().remove(pe)
    sp_after=OxmlElement('w:p'); sppr=OxmlElement('w:pPr'); spc=OxmlElement('w:spacing'); spc.set(qn('w:after'),'120'); sppr.append(spc); sp_after.append(sppr)
    tbl.addnext(sp_after)

def style_callouts():
    # Group a callout's FULL contiguous block: the first Callout*-styled paragraph,
    # then every following paragraph that is the SAME callout style, a bullet, or
    # empty — up to and including the LAST callout-styled paragraph in the run.
    # (Pandoc tags bullets inside the div as Normal+numPr, not the callout style,
    # which previously split the box. This consumes them into one box.)
    children = list(body.iterchildren())
    i = 0
    groups = []
    while i < len(children):
        ch = children[i]
        st = _para_style(ch)
        if st in CALLOUT_STYLES:
            grp=[ch]; j=i+1; last_callout=0  # index within grp of last callout-styled p
            while j < len(children):
                cj = children[j]
                sj = _para_style(cj)
                if sj == st:
                    grp.append(cj); last_callout=len(grp)-1; j+=1
                elif _is_bullet(cj) or _is_empty_p(cj):
                    grp.append(cj); j+=1
                else:
                    break
            # trim anything after the last callout-styled paragraph (don't swallow
            # trailing bullets/blanks that belong to the next section)
            grp = grp[:last_callout+1]
            groups.append((st, grp)); i = i + len(grp)
        else:
            i+=1
    for st, grp in groups:
        fill,bar,lblcol = CALLOUT_STYLES[st]
        _wrap_group_in_box(grp, fill=fill, bar=bar, labelcolor=lblcol)

def box_paragraph(p, *, fill, bar, textcolor, italic=False):
    """Single-paragraph shaded box (chapter abstract)."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')): pPr.remove(old)
    shade(pPr, fill)
    pbdr = OxmlElement('w:pBdr')
    for edge,(sz,color) in (('left',('24',bar)),('top',('6',fill)),('bottom',('6',fill)),('right',('6',fill))):
        b=OxmlElement(f'w:{edge}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz); b.set(qn('w:space'),'8'); b.set(qn('w:color'),color); pbdr.append(b)
    pPr.append(pbdr)
    ind=OxmlElement('w:ind'); ind.set(qn('w:left'),'180'); ind.set(qn('w:right'),'120'); pPr.append(ind)
    sp=OxmlElement('w:spacing'); sp.set(qn('w:before'),'120'); sp.set(qn('w:after'),'160'); pPr.append(sp)
    for r in p.runs:
        if textcolor: r.font.color.rgb = RGBColor.from_string(textcolor)
        if italic: r.italic=True

def _para_style(p_el):
    if p_el.tag != qn('w:p'): return None
    pPr=p_el.find(qn('w:pPr'))
    if pPr is None: return None
    ps=pPr.find(qn('w:pStyle'))
    return ps.get(qn('w:val')) if ps is not None else None

def _wrap_group_in_box(p_elements, *, fill, bar, labelcolor):
    """Move a group of paragraph elements into one single-cell shaded table."""
    first = p_elements[0]
    # build table
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(PORTRAIT_DXA)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    # left bar via table borders (thick left, none others)
    bd=OxmlElement('w:tblBorders')
    for edge,(sz,color) in (('left',('24',bar)),('top',('2',fill)),('bottom',('2',fill)),
                            ('right',('2',fill)),('insideH',('2',fill)),('insideV',('2',fill))):
        b=OxmlElement(f'w:{edge}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz); b.set(qn('w:space'),'0'); b.set(qn('w:color'),color); bd.append(b)
    tblPr.append(bd)
    cm=OxmlElement('w:tblCellMar')
    for e,w in (('top','120'),('left','160'),('bottom','120'),('right','160')):
        m=OxmlElement(f'w:{e}'); m.set(qn('w:w'),w); m.set(qn('w:type'),'dxa'); cm.append(m)
    tblPr.append(cm)
    tbl.append(tblPr)
    grid=OxmlElement('w:tblGrid'); gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(PORTRAIT_DXA)); grid.append(gc); tbl.append(grid)
    tr=OxmlElement('w:tr'); tc=OxmlElement('w:tc')
    tcPr=OxmlElement('w:tcPr'); tcW=OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(PORTRAIT_DXA)); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
    tc.append(tcPr)
    # move paragraphs into the cell, restyling them
    first.addprevious(tbl)
    for k,pe in enumerate(p_elements):
        pe.getparent().remove(pe)
        # strip the Callout* pStyle -> Normal, set spacing tight
        pPr = pe.find(qn('w:pPr'))
        if pPr is None: pPr=OxmlElement('w:pPr'); pe.insert(0,pPr)
        ps=pPr.find(qn('w:pStyle'))
        if ps is not None: pPr.remove(ps)
        for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
        sp=OxmlElement('w:spacing'); sp.set(qn('w:after'),'60'); sp.set(qn('w:line'),'264'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
        # body runs: smaller (10pt) + italic; label (first para, bold run) colored, NOT italic
        is_label = (k==0)
        for r in pe.findall(qn('w:r')):
            rpr=r.find(qn('w:rPr'))
            if rpr is None: rpr=OxmlElement('w:rPr'); r.insert(0,rpr)
            # size 20 half-pts = 10pt
            for tag in ('w:sz','w:szCs'):
                e=rpr.find(qn(tag))
                if e is None: e=OxmlElement(tag); rpr.append(e)
                e.set(qn('w:val'),'20')
            bold = rpr.find(qn('w:b')) is not None
            if is_label and bold:
                col=rpr.find(qn('w:color'))
                if col is None: col=OxmlElement('w:color'); rpr.append(col)
                col.set(qn('w:val'),labelcolor)
                # label slightly bigger
                for tag in ('w:sz','w:szCs'):
                    rpr.find(qn(tag)).set(qn('w:val'),'21')
            # body stays roman (NOT italic) — smaller size already differentiates it
        tc.append(pe)
    tr.append(tc); tbl.append(tr)
    # spacer paragraph after box
    sp_after=OxmlElement('w:p'); sppr=OxmlElement('w:pPr'); spc=OxmlElement('w:spacing'); spc.set(qn('w:after'),'120'); sppr.append(spc); sp_after.append(sppr)
    tbl.addnext(sp_after)

# ── 4. Sources -> small font ─────────────────────────────────────────────────
def _sources_toprule(p):
    """Thin gray rule above a Sources block so it reads as a separated end-of-section
    footer, with generous space before it (pushed down from the body text above)."""
    pPr=p._p.get_or_add_pPr()
    pbdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:top'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
    b.set(qn('w:space'),'8'); b.set(qn('w:color'),'B8B8B8'); pbdr.append(b)
    pPr.append(pbdr)
    sp=pPr.find(qn('w:spacing'))
    if sp is None: sp=OxmlElement('w:spacing'); pPr.append(sp)
    sp.set(qn('w:before'),'360'); sp.set(qn('w:after'),'0')  # push down from body above

def shrink_sources():
    """Normalize every Sources block to ONE consistent look regardless of whether the
    manuscript wrote it as a '## Sources' heading or an inline 'Sources: ...' line:
    small 9pt italic gray citation text, a bold 'Sources' lead-in, and a thin top
    rule with space-before so it sits apart as an end-of-section footer."""
    GRAY=RGBColor(0x40,0x40,0x40)
    def _is_sources_heading(p):
        # strip only a leading section-number prefix like "1.18 " (digits/dots + space),
        # NOT a leading capital letter — the old regex ate the 'S' of 'Sources'.
        t=re.sub(r'^[0-9]+(\.[0-9A-Z]+)*\s+','',p.text.strip())
        return p.style.name.startswith('Heading') and t=='Sources'

    # PASS 1: demote EVERY 'Sources' heading to a small bold Normal label (keeps it
    # out of the TOC). Snapshot the list first so restyling doesn't disturb iteration.
    heading_ids=[]
    for p in list(doc.paragraphs):
        if _is_sources_heading(p):
            p.style=doc.styles['Normal']
            if p.runs:
                r0=p.runs[0]
                # The run carries an rPr inherited from the Heading style (navy
                # 2E74B5, 10pt). Assigning .font.* leaves those direct-formatting
                # elements in place, so the label kept rendering as a heading —
                # which is exactly what made 3 of 19 Sources blocks look wrong.
                # Clear the run properties first, then restyle from scratch.
                old=r0._r.find(qn('w:rPr'))
                if old is not None: r0._r.remove(old)
                r0.text='Sources'; r0.bold=True; r0.italic=False
                r0.font.size=Pt(9); r0.font.color.rgb=GRAY
                for extra in p.runs[1:]: extra.text=''
            # paragraph-level marks (heading colour/size/keepNext) also survive
            # the style swap — drop them so Normal actually governs the label
            pPr=p._p.find(qn('w:pPr'))
            if pPr is not None:
                for tag in ('w:rPr','w:keepNext','w:keepLines','w:outlineLvl'):
                    for el in pPr.findall(qn(tag)): pPr.remove(el)
            _sources_toprule(p)
            heading_ids.append(id(p._p))

    # PASS 2: style the citation lines that follow each demoted 'Sources' label
    paras=doc.paragraphs
    for i,p in enumerate(paras):
        if id(p._p) in heading_ids:
            j=i+1
            while j<len(paras) and not paras[j].style.name.startswith('Heading') and paras[j].text.strip():
                for r in paras[j].runs:
                    r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GRAY
                j+=1

def finalize_sources():
    """FINAL, AUTHORITATIVE styling for every 'Sources:' paragraph.

    Runs LAST in the pipeline so it overwrites whatever earlier passes (and
    pandoc's own Calibri run properties) produced. A citation block then looks
    identical whether the manuscript wrote it as a '## Sources' heading or an
    inline 'Sources: ...' line — the two used to drift apart, leaving three
    chapters in Calibri/black with different spacing from the other fifteen.
    """
    for p in doc.paragraphs:
        if not p.text.strip().startswith('Sources:'):
            continue
        p.style = doc.styles['Body Text']
        pf = p.paragraph_format
        pf.space_before = Pt(18); pf.space_after = Pt(6)
        for r in p.runs:
            old = r._r.find(qn('w:rPr'))
            if old is not None: r._r.remove(old)
            r.bold = False; r.italic = True
            r.font.size = Pt(9); r.font.color.rgb = GRAY
            r.font.name = BODY_FONT
            rPr = r._r.get_or_add_rPr()
            rf = OxmlElement('w:rFonts')
            for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), BODY_FONT)
            rPr.insert(0, rf)
        _sources_toprule(p)


# ── 5. tables: full width portrait; wide -> smaller font ─────────────────────
def style_tables():
    for t in doc.tables:
        # skip callout boxes (single-cell tables we built earlier)
        if len(t.rows)==1 and len(t.columns)==1:
            continue
        # skip stat strips (already fully styled by _build_stat_cards)
        capel=t._tbl.tblPr.find(qn('w:tblCaption')) if t._tbl.tblPr is not None else None
        if capel is not None and capel.get(qn('w:val'))=='STATSTRIP':
            continue
        try: t.style='HTR Table'
        except Exception: pass
        tblPr=t._tbl.tblPr
        look=tblPr.find(qn('w:tblLook'))
        if look is None: look=OxmlElement('w:tblLook'); tblPr.append(look)
        look.set(qn('w:firstRow'),'1'); look.set(qn('w:noHBand'),'0'); look.set(qn('w:noVBand'),'1')
        # navy header fill manually (in case style not honored)
        hdr=t.rows[0]
        for c in hdr.cells:
            tcPr=c._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
            shade(tcPr,'1B3A6B')
            for para in c.paragraphs:
                for r in para.runs: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.bold=True
        # full width, but AUTOFIT columns to content (not equal width)
        for tag in ('w:tblW','w:tblLayout'):
            el=tblPr.find(qn(tag))
            if el is not None: tblPr.remove(el)
        tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(PORTRAIT_DXA)); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
        lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'autofit'); tblPr.append(lay)
        ncol=len(t.columns)
        # size columns by max content length (so a 1-char "#" col stays narrow)
        weights=[]
        for ci in range(ncol):
            maxlen=1
            for row in t.rows:
                try: maxlen=max(maxlen, len(row.cells[ci].text.strip()))
                except IndexError: pass
            weights.append(min(maxlen, 60))  # cap so one huge cell doesn't dominate
        total=sum(weights) or 1
        widths=[max(int(PORTRAIT_DXA*w/total), 360) for w in weights]  # min ~0.25in
        # normalize to exactly PORTRAIT_DXA
        scale=PORTRAIT_DXA/sum(widths); widths=[int(x*scale) for x in widths]
        grid=t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            gcs=grid.findall(qn('w:gridCol'))
            for gc,w in zip(gcs,widths): gc.set(qn('w:w'),str(w))
        for row in t.rows:
            for ci,cell in enumerate(row.cells):
                if ci<len(widths):
                    tcPr=cell._tc.get_or_add_tcPr()
                    tcW=tcPr.find(qn('w:tcW'))
                    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
                    tcW.set(qn('w:w'),str(widths[ci])); tcW.set(qn('w:type'),'dxa')
        # Table text must be SMALLER than body (body is 10.5pt). Force 9pt on every
        # cell run (overrides Pandoc's explicit sizes); wide+dense tables go smaller
        # still (8pt) since they also rotate to landscape.
        cell_pt = Pt(8) if _table_is_wide_dense(t) else Pt(9)
        for row in t.rows:
            for c in row.cells:
                for para in c.paragraphs:
                    for r in para.runs:
                        r.font.size=cell_pt
        # allow rows to split across pages so a tall table fills the page
        # instead of jumping whole to the next page (the main cause of big voids)
        for row in t.rows:
            trPr = row._tr.get_or_add_trPr()
            for old in trPr.findall(qn('w:cantSplit')): trPr.remove(old)
        # NOTE: no post-table spacer paragraph — it separated tables from their
        # captions. Caption style (space-after) + body spacing handle the gap.

def style_dependency_matrix():
    """Make the Figure 1.B six-pillar dependency matrix legible and clean:
      - blank the '—' filler cells (they read as spreadsheet noise);
      - shade the diagonal (self->self) cells light gray so it reads as 'N/A';
      - tint the LEFT header column navy like the top header row (both axes labeled);
      - color-code each populated cell's leading verb by dependency type, bold;
      - center every cell, comfortable padding.
    Runs after style_tables(); the matrix is forced landscape by _table_is_wide_dense."""
    TYPE_COLOR = {'Enables':RGBColor(0xC9,0x7A,0x0F), 'Requires':RGBColor(0x1F,0x6B,0x5D),
                  'Drives':RGBColor(0x3E,0x4F,0xB0), 'Constrains':RGBColor(0x9A,0x2A,0x48)}
    for t in doc.tables:
        if not _is_dependency_matrix(t):
            continue
        ncol = len(t.columns)

        # ── purpose-built column widths (PORTRAIT) ───────────────────────────
        # First column holds the pillar names on one line; the six pillar columns
        # split the remaining portrait width evenly. (Author does final matrix
        # formatting manually — this is a reasonable starting point.)
        TOTAL_W = PORTRAIT_DXA
        first_w = int(0.95*1440)                       # pillar names on one line
        rest_w  = (TOTAL_W - first_w)//(ncol-1)        # even split of the remainder
        widths  = [first_w] + [rest_w]*(ncol-1)
        widths[-1] += TOTAL_W - sum(widths)            # normalize to exact total
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
                gc.set(qn('w:w'), str(w))
        # fixed layout so Word HONORS these widths (autofit would re-shrink col 1)
        tblPr = t._tbl.tblPr
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
                tcW.set(qn('w:w'), str(widths[ci])); tcW.set(qn('w:type'),'dxa')
                # never split a WORD inside any cell (break at spaces only)
                for old in tcPr.findall(qn('w:noWrap')): tcPr.remove(old)

        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                tcPr = cell._tc.get_or_add_tcPr()
                # center vertically
                for old in tcPr.findall(qn('w:vAlign')): tcPr.remove(old)
                va = OxmlElement('w:vAlign'); va.set(qn('w:val'),'center'); tcPr.append(va)
                # LEFT header column (row labels): navy like the top row
                if ci == 0 and ri > 0:
                    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
                    shade(tcPr,'1B3A6B')
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pPr = para._p.get_or_add_pPr()
                        if pPr.find(qn('w:suppressAutoHyphens')) is None:
                            pPr.append(OxmlElement('w:suppressAutoHyphens'))
                        for r in para.runs:
                            r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.bold=True; r.font.size=Pt(8)
                    continue
                if ri == 0:  # top header row already navy from style_tables
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in para.runs:
                            r.font.size=Pt(8)
                    continue
                # body cells
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if txt in ('—','-',''):
                    # diagonal / empty: blank + faint gray fill
                    for r in list(para.runs): r.text=''
                    if ri == ci:  # diagonal self-cell
                        for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
                        shade(tcPr,'ECEFF3')
                    continue
                # populated cell: bold + color the leading verb
                verb = txt.split(' ',1)[0]
                col = TYPE_COLOR.get(verb)
                for r in para.runs:
                    r.font.size=Pt(8)
                    if col is not None:
                        r.font.color.rgb=col; r.bold=True

def _sectPr(landscape):
    """Build a <w:sectPr>. Portrait keeps 1in margins. Landscape uses TIGHTER
    margins (0.5in top/bottom, 0.7in sides) so a dense table that would otherwise
    spill one row onto a second page fits on a single landscape page."""
    sp = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    if landscape:
        pgSz.set(qn('w:w'), '15840'); pgSz.set(qn('w:h'), '12240'); pgSz.set(qn('w:orient'), 'landscape')
        margins = (('top','720'),('bottom','720'),('left','1008'),('right','1008'),
                   ('header','360'),('footer','360'),('gutter','0'))
    else:
        pgSz.set(qn('w:w'), '12240'); pgSz.set(qn('w:h'), '15840'); pgSz.set(qn('w:orient'), 'portrait')
        # book margins: 0.9in top/bottom (1296), 1.1in sides (1584) — match reference.docx
        margins = (('top','1296'),('bottom','1296'),('left','1584'),('right','1584'),
                   ('header','648'),('footer','648'),('gutter','0'))
    sp.append(pgSz)
    mar = OxmlElement('w:pgMar')
    for k, v in margins:
        mar.set(qn('w:'+k), v)
    sp.append(mar)
    return sp

def wide_tables_to_landscape():
    """Put each wide table (>= WIDE_TABLE_MIN_COLS columns) on its OWN landscape
    page, isolated so the rotation cannot bleed into surrounding content.

    CRITICAL section-break rule (this is what broke it before): in Word a sectPr
    placed inside a paragraph applies BACKWARD — it ends the section that PRECEDES
    it. So to make ONLY the table landscape:
      - put a PORTRAIT sectPr paragraph BEFORE the table  -> closes the preceding
        content as portrait (cover/TOC/prose stay upright);
      - put a LANDSCAPE sectPr paragraph AFTER the table+caption -> the section
        that ENDS there (i.e. the table) is landscape;
      - the document's trailing default sectPr stays portrait -> the tail is upright.

    Also re-widens the table + its caption to the landscape text width.
    """
    for t in list(doc.tables):
        if not _table_is_wide_dense(t):
            continue
        tbl_el = t._tbl

        # 1) re-widen table + columns to landscape width
        tblPr = tbl_el.tblPr
        w = tblPr.find(qn('w:tblW'))
        if w is None:
            w = OxmlElement('w:tblW'); tblPr.append(w)
        w.set(qn('w:w'), str(LANDSCAPE_DXA)); w.set(qn('w:type'), 'dxa')
        grid = tbl_el.find(qn('w:tblGrid'))
        if grid is not None:
            gcs = grid.findall(qn('w:gridCol'))
            # scale existing (content-weighted) widths up to the landscape total
            cur = [int(gc.get(qn('w:w')) or 0) for gc in gcs]
            tot = sum(cur) or 1
            for gc, cw in zip(gcs, cur):
                gc.set(qn('w:w'), str(int(LANDSCAPE_DXA * cw / tot)))

        # 2) PORTRAIT sectPr paragraph BEFORE the table -> closes the preceding
        #    content as a portrait section (everything up to here stays upright).
        p_before = OxmlElement('w:p')
        pPr_b = OxmlElement('w:pPr'); pPr_b.append(_sectPr(landscape=False)); p_before.append(pPr_b)
        tbl_el.addprevious(p_before)

        # 3) LANDSCAPE sectPr paragraph AFTER the table (+ its caption) -> the
        #    section that ENDS at this marker is the table, so the table is
        #    landscape. Content after this marker belongs to the next section,
        #    governed by the document's trailing (portrait) default sectPr.
        after_anchor = tbl_el
        nxt = tbl_el.getnext()
        if nxt is not None and nxt.tag == qn('w:p'):
            txt = ''.join(x.text or '' for x in nxt.iter(qn('w:t'))).strip()
            if re.match(r'^(Figure|Table)\b', txt):
                after_anchor = nxt  # keep the caption on the landscape page too
        p_after = OxmlElement('w:p')
        pPr_a = OxmlElement('w:pPr'); pPr_a.append(_sectPr(landscape=True)); p_after.append(pPr_a)
        after_anchor.addnext(p_after)

        # 4) BACKFILL the white space on the page BEFORE the landscape table:
        #    move the body paragraphs that flow AFTER the table up to BEFORE the
        #    portrait break, so they fill the otherwise-empty portion of the prior
        #    page and the table starts cleanly on the next (landscape) page.
        #    Only move plain body prose — stop at the next heading, table, or an
        #    already-placed section break, and cap how much we relocate.
        def _is_heading(el):
            if el is None or el.tag != qn('w:p'): return False
            st = el.find(qn('w:pPr'))
            if st is None: return False
            ps = st.find(qn('w:pStyle'))
            return ps is not None and (ps.get(qn('w:val')) or '').startswith(('Heading','Title'))
        def _has_sectPr(el):
            return el is not None and el.find('.//'+qn('w:sectPr')) is not None
        moved = 0
        cursor = p_after.getnext()
        while cursor is not None and moved < 6:            # cap: at most 6 paragraphs
            nxt2 = cursor.getnext()
            # stop conditions: next heading, a table, or another section break
            if cursor.tag == qn('w:tbl') or _is_heading(cursor) or _has_sectPr(cursor):
                break
            if cursor.tag == qn('w:p'):
                txt = ''.join(x.text or '' for x in cursor.iter(qn('w:t'))).strip()
                if txt == '':                              # skip/absorb empty spacers
                    cursor = nxt2; continue
                # relocate this body paragraph to just before the portrait break
                cursor.getparent().remove(cursor)
                p_before.addprevious(cursor)
                moved += 1
            cursor = nxt2

def style_captions():
    """Apply the Caption style to Figure/Table caption paragraphs so they hug the
    table above and sit well clear of the content below."""
    import re as _re
    for p in doc.paragraphs:
        if _re.match(r'^(Figure|Table)\s+[\w.]+\s*[—(-]', p.text.strip()):
            try: p.style = doc.styles['Caption']
            except Exception: pass

def style_inline_figures():
    """Size + center inline body images that came from markdown ![](...) syntax.
    Pandoc drops them in at native pixel size, left-aligned; we cap the width to
    the text column and center the paragraph. The cover image is inserted later by
    add_cover() and is not present in the body yet, so it is untouched here."""
    TEXT_W = Inches(6.0)  # sit comfortably within the 6.5in text column
    for p in doc.paragraphs:
        blips = p._p.findall('.//' + qn('a:blip'))
        if not blips:
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            for inline in run._r.findall('.//' + qn('wp:inline')) + run._r.findall('.//' + qn('wp:anchor')):
                extent = inline.find(qn('wp:extent'))
                if extent is None:
                    continue
                cx = int(extent.get('cx')); cy = int(extent.get('cy'))
                if cx <= TEXT_W:
                    continue
                ratio = TEXT_W / cx
                new_cx = int(TEXT_W); new_cy = int(cy * ratio)
                extent.set('cx', str(new_cx)); extent.set('cy', str(new_cy))
                # also update the inner picture extent (a:ext) so it scales, not crops
                for ext in inline.findall('.//' + qn('a:ext')):
                    ext.set('cx', str(new_cx)); ext.set('cy', str(new_cy))

# ── 6. cover + TOC own page + chapter page breaks ────────────────────────────
def page_break_para():
    p=OxmlElement('w:p'); r=OxmlElement('w:r'); br=OxmlElement('w:br'); br.set(qn('w:type'),'page')
    r.append(br); p.append(r); return p

def _bottom_rule(para, color='1B3A6B', sz='18'):
    """Add a navy bottom border to a paragraph (the rule under the title)."""
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),sz)
    bottom.set(qn('w:space'),'6'); bottom.set(qn('w:color'),color)
    pbdr.append(bottom); pPr.append(pbdr)

def _set_runs(para, *, size, color, bold=False, italic=False, caps=False):
    NAVY = RGBColor(0x1B,0x3A,0x6B)
    for r in para.runs:
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = bold; r.italic = italic
        if caps: r.text = r.text.upper()

def add_cover():
    """Style the cover title block to match the author's reference layout, then
    place the cover image centered below it and close the page:
        TRANSFORMING AMERICAN HEALTHCARE   (navy bold ~26pt, uppercase)
        ───────────────────────────────    (navy rule)
        A Six-Pillar Framework ...          (navy italic subtitle)
        With Vermont as the National ...    (navy italic, smaller)
        Bechir BenSaid                      (bold, author)
        April 2026                          (gray, small)
        [cover image, centered]
    """
    NAVY = RGBColor(0x1B,0x3A,0x6B)
    GRAY = RGBColor(0x55,0x55,0x55)
    BLACK = RGBColor(0x1A,0x1A,0x1A)

    # locate the five leading title-block paragraphs by content
    title=sub1=sub2=author=edition=None
    for p in doc.paragraphs[:12]:
        tl = p.text.strip()
        low = tl.lower()
        if not tl: continue
        if 'transforming american healthcare' in low and title is None: title=p
        elif low.startswith('a six-pillar') and sub1 is None: sub1=p
        elif low.startswith('with vermont') and sub2 is None: sub2=p
        elif ('bechir' in low or low.startswith('health transformation review')) and author is None: author=p
        elif ('edition' in low or 'april 2026' in low) and edition is None: edition=p

    if title is not None:
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_runs(title, size=26, color=NAVY, bold=True, caps=True)
        title.paragraph_format.space_before = Pt(72)   # push down from top edge
        title.paragraph_format.space_after = Pt(6)
        _bottom_rule(title)                            # navy rule under the title
    if sub1 is not None:
        sub1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_runs(sub1, size=14, color=NAVY, italic=True)
        sub1.paragraph_format.space_before = Pt(10); sub1.paragraph_format.space_after = Pt(2)
    if sub2 is not None:
        sub2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_runs(sub2, size=11.5, color=NAVY, italic=True)
        sub2.paragraph_format.space_after = Pt(24)
    if author is not None:
        # Style only — never rewrite the byline. The .md is the source of truth
        # for cover text; an override here would silently undo any change the
        # author makes to it (same bug as the hard-coded date below).
        author.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_runs(author, size=12, color=BLACK, bold=True)
        author.paragraph_format.space_before = Pt(4); author.paragraph_format.space_after = Pt(2)
    if edition is not None:
        # Style it, but do NOT rewrite its text. This used to hard-code
        # 'April 2026', which silently restored the date every build after the
        # author had deleted it from the manuscript (2026-07-27). The cover
        # line says whatever the .md says; if the .md has no date line, there
        # is no `edition` paragraph to style and nothing is printed.
        edition.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_runs(edition, size=10, color=GRAY)
        edition.paragraph_format.space_after = Pt(0)

    if not (COVER and os.path.exists(COVER)):
        # no image: still drop a page break after the title block
        anchor = edition or author or sub2 or title or doc.paragraphs[0]
    else:
        anchor = edition or author or sub2 or title or doc.paragraphs[0]

    def _new_para_after(ref):
        new = OxmlElement('w:p'); ref._p.addnext(new)
        from docx.text.paragraph import Paragraph
        return Paragraph(new, ref._parent)

    # spacer -> centered image -> page break
    sp = _new_para_after(anchor)
    sp.paragraph_format.space_after = Pt(48)   # breathing room, image lower on page

    last = sp
    if COVER and os.path.exists(COVER):
        img_p = _new_para_after(sp)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(COVER, width=Inches(5.6))
        last = img_p

    brk = _new_para_after(last)
    r = brk.add_run(); br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r._r.append(br)

def first_h1_el():
    for ch in body.iterchildren():
        if ch.tag==qn('w:p'):
            pPr=ch.find(qn('w:pPr'))
            if pPr is not None:
                ps=pPr.find(qn('w:pStyle'))
                if ps is not None and ps.get(qn('w:val')) in ('Heading1','Heading 1'):
                    return ch
    return None

def add_page_numbers():
    """Bottom-right page numbers that START AT 1 ON THE PREFACE.

    The cover and the Table of Contents must carry no number, so the document is
    split into two sections:
      - section 1 = cover + TOC          -> its own footer, left empty
      - section 2 = Preface -> end       -> footer with a PAGE field, numbering
                                            restarted at 1

    The split is made by giving the LAST paragraph before the Preface (the TOC's
    trailing page-break paragraph) its own sectPr. In Word a sectPr on a
    paragraph ends the section AT that paragraph, so everything after it belongs
    to the next section — the same rule the landscape-table code relies on.
    """
    h1 = first_h1_el()                      # the Preface heading
    if h1 is None: return

    # ---- 1. build the two footers on the document's single sectPr ------------
    # python-docx exposes section objects; the document currently has one.
    sec_all = doc.sections[0]

    def _footer_page_field(footer, with_number):
        """Right-aligned footer paragraph; PAGE field only when with_number."""
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        if not with_number:
            return
        r = p.add_run()
        fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
        for el in (fld_b, instr, fld_e):
            r._element.append(el)
        _set_runs(p, size=9.5, color=RGBColor(0x55, 0x55, 0x55))

    # ---- 2. create the numbered footer FIRST -------------------------------
    # This must happen before the deepcopy below: touching sec_all.footer is what
    # materialises footer1.xml and adds the <w:footerReference> to the document
    # sectPr. Copying first would clone a sectPr that has no footer yet, and the
    # reference would then land on the wrong (front-matter) section.
    _footer_page_field(sec_all.footer, True)
    sec_all.different_first_page_header_footer = False

    # restart numbering at 1 on the body section (the document default sectPr,
    # which governs everything from the last section break to the end)
    pgnum = sec_all._sectPr.find(qn('w:pgNumType'))
    if pgnum is None:
        pgnum = OxmlElement('w:pgNumType')
        sec_all._sectPr.append(pgnum)
    pgnum.set(qn('w:start'), '1')

    prev = h1.getprevious()
    if prev is None or prev.tag != qn('w:p'):
        return                       # nothing to split on: number everything

    # ---- 3. split: the paragraph before the Preface ends the front section ---
    # A paragraph-level sectPr terminates the section AT that paragraph, so this
    # one governs the cover + TOC. It inherits the page setup but must carry
    # neither the footer nor the restart, or the front matter would be numbered.
    front_sectPr = copy.deepcopy(sec_all._sectPr)
    for ref in front_sectPr.findall(qn('w:footerReference')):
        front_sectPr.remove(ref)
    for pn in front_sectPr.findall(qn('w:pgNumType')):
        front_sectPr.remove(pn)

    # Suppress numbering on the front matter EXPLICITLY rather than by omission.
    # Google Docs, on export, fills in a <w:pgNumType w:start="1"/> wherever the
    # element is simply absent — which restarted the count on the cover page and
    # made the Preface land on 3 instead of 1. Writing an explicit start here,
    # plus an empty footer reference below, leaves Docs nothing to "helpfully"
    # complete.
    titlePg = OxmlElement('w:titlePg')
    front_sectPr.append(titlePg)

    pPr = prev.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); prev.insert(0, pPr)
    pPr.append(front_sectPr)

    # ---- 4. every intermediate section (landscape tables etc.) must inherit ---
    # the numbered footer. A sectPr with no footerReference of its own inherits
    # from the PREVIOUS section, and the landscape blocks sit between the front
    # matter and the body — without an explicit reference the first of them
    # would inherit the footerless front-matter section and blank the numbers
    # from there on. Point them all at the same footer part.
    fref = sec_all._sectPr.find(qn('w:footerReference'))
    if fref is None:
        return
    rid = fref.get(qn('r:id'))
    for sp in body.iter(qn('w:sectPr')):
        if sp is sec_all._sectPr or sp is front_sectPr:
            continue
        if sp.find(qn('w:footerReference')) is not None:
            continue
        ref = OxmlElement('w:footerReference')
        ref.set(qn('w:type'), 'default'); ref.set(qn('r:id'), rid)
        sp.insert(0, ref)

def _toc_line(text, name, level):
    """One TOC entry paragraph: styled TOC 1/2/3, the heading text as a hyperlink
    to its _Toc bookmark, a right dot-leader tab, then a PAGEREF field that shows
    the real page number. PAGEREF is computed by Word AND Google Docs on refresh,
    so numbers are correct in both — unlike a bare TOC field, which GDocs left at 1."""
    p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
    ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), {1:'TOC1',2:'TOC2',3:'TOC3'}[level]); pPr.append(ps)
    p.append(pPr)
    # hyperlink wrapping the heading text
    hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), name)
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = text
    r.append(t); hl.append(r); p.append(hl)
    # tab to the right dot-leader stop
    rt = OxmlElement('w:r'); tab = OxmlElement('w:tab'); rt.append(tab); p.append(rt)
    # PAGEREF field:  { PAGEREF _Tocxxxx \h }
    def _run_with(child):
        rr = OxmlElement('w:r'); rr.append(child); return rr
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve')
    instr.text = f' PAGEREF {name} \\h '
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate')
    ph = OxmlElement('w:r'); pht = OxmlElement('w:t'); pht.text = '1'; ph.append(pht)
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    for el in (_run_with(b), _run_with(instr), _run_with(sep), ph, _run_with(end)):
        p.append(el)
    return p

def insert_toc():
    """Build the table of contents as explicit TOC-styled entries, each carrying a
    live PAGEREF field to its heading's _Toc bookmark. Page numbers with dot
    leaders update automatically on refresh in BOTH Word and Google Docs. The
    whole block is wrapped in a real TOC field so Word/GDocs still treat it as a
    table of contents (right-click > Update, or GDocs refresh)."""
    h1 = first_h1_el()
    if h1 is None: return

    # label paragraph (Title style so it reads as a page header)
    lp = OxmlElement('w:p'); lpr = OxmlElement('w:pPr'); lps = OxmlElement('w:pStyle')
    lps.set(qn('w:val'),'Title'); lpr.append(lps); lp.append(lpr)
    lr = OxmlElement('w:r'); lrpr = OxmlElement('w:rPr'); lc = OxmlElement('w:color'); lc.set(qn('w:val'),'1B3A6B')
    lrpr.append(lc); lr.append(lrpr); lt = OxmlElement('w:t'); lt.text = 'Table of Contents'; lr.append(lt); lp.append(lr)

    lvlmap = {'Heading 1':1,'Heading 2':2}   # 2-level TOC: chapters + sections only
    entries = [(lvlmap[st], txt, name) for (st, txt, name) in _HEADING_BOOKMARKS
               if st in lvlmap and txt != 'Table of Contents']

    # open a TOC field, drop the explicit entries as its result, then close it.
    def _run_with(child):
        r = OxmlElement('w:r'); r.append(child); return r
    open_p = OxmlElement('w:p')
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate')
    for el in (_run_with(b), _run_with(instr), _run_with(sep)):
        open_p.append(el)

    entry_ps = [_toc_line(txt, name, lvl) for (lvl, txt, name) in entries]

    close_p = entry_ps[-1] if entry_ps else open_p
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    close_p.append(_run_with(end))

    nodes = [page_break_para(), lp, open_p] + entry_ps + [page_break_para()]
    for el in nodes:
        h1.addprevious(el)

def chapter_page_breaks():
    first=True
    for p in list(doc.paragraphs):
        if p.style.name=='Heading 1' and (p.text.strip().startswith('Chapter') or p.text.strip().startswith('Appendix') or p.text.strip() in ('PREFACE','INTRODUCTION','Conclusion')):
            if first: first=False; continue
            p._p.addprevious(page_break_para())

def rule_above_major_sections():
    """Draw a thin rule ABOVE every Heading 2 so a new major section is
    unmistakable.

    The reported problem: a long run of bolded items (the counterargument
    sections, the analytical-payoff list) ends and a brand-new major section
    begins, but nothing tells the eye that the list is over — the new section
    reads as one more item in it. Size alone did not carry that signal.

    A hairline above the heading, plus the widened space-before in the
    stylesheet, makes the boundary explicit. Chapter openers (Heading 1) already
    start a fresh page and need no rule.
    """
    for p in doc.paragraphs:
        if p.style is None or p.style.name not in ('Heading 2',):
            continue
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn('w:pBdr')) is not None:      # never double-rule
            continue
        # a heading sitting at the very top of a page would show a floating line
        if pPr.find(qn('w:pageBreakBefore')) is not None:
            continue
        pbdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')      # hairline
        top.set(qn('w:space'), '10')                                  # gap heading↔rule
        top.set(qn('w:color'), 'C9D2DF')                              # pale navy
        pbdr.append(top); pPr.append(pbdr)

def tighten_table_pagination():
    """Tables flow. No exceptions.

    THE RULE (author's, 2026-07-27): every table may split across a page break,
    and every table repeats its header row at the top of each new page. Nothing
    is held back to keep a table whole.

    Why this and not something cleverer: a table that cannot split forces Word
    to shove the whole thing to the next page, leaving a half-blank page behind
    it. Earlier versions tried to pick which tables were "worth" keeping whole —
    first by row count, then by estimated height — and every threshold left some
    tables locked and some pages half empty. Letting them all flow removes the
    cause instead of tuning the symptom. A split table is not a defect when the
    header repeats: the reader still sees what every column means.

    Also releases keepNext on the paragraphs above each table, which otherwise
    drags the lead-in text forward onto the table's page and widens the very gap
    it was trying to avoid. The lead-ins name their figure explicitly
    ("Figure 1.A characterizes…"), so they read fine a page earlier.
    """
    for tbl in doc.tables:
        rows = tbl.rows
        if not rows:
            continue

        # 1. every row may break across pages
        for r in rows:
            trPr = r._tr.get_or_add_trPr()
            for old in trPr.findall(qn('w:cantSplit')):
                trPr.remove(old)

        # 2. the header row repeats on every page the table continues onto
        trPr = rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true')
            trPr.append(el)
        # a repeating header must not itself be split mid-row
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))

        # 3. let the text above the table stay where it is.
        #
        # This includes HEADINGS. A heading carries keepNext from its style, so
        # Word refuses to separate it from the table below — and when the table
        # will not fit, it drags the heading forward too, stranding everything
        # above them and leaving a half-blank page. Twenty tables in this book
        # sit directly under a heading, and they were the largest single source
        # of that white space. Releasing keepNext lets the heading sit at the
        # bottom of the page with the table beginning right after it, or flow
        # onto the next page on its own — either way the preceding page fills.
        # FIRST: delete empty paragraphs sitting between the lead-in and the
        # table. Pandoc turns the blank line that markdown requires before a
        # table into a real empty paragraph. It renders as a blank line, it
        # breaks the keepNext chain from the lead-in to the table, and it is
        # the reason "Figure 1.A characterizes…" kept getting stranded at the
        # top of a page with the table overleaf. 49 tables had one.
        while True:
            prev_el = tbl._tbl.getprevious()
            if (prev_el is not None and prev_el.tag == qn('w:p')
                    and not "".join(prev_el.itertext()).strip()
                    and prev_el.find(qn('w:drawing')) is None):
                prev_el.getparent().remove(prev_el)
                continue
            break

        # EXCEPTION: a figure lead-in ("Figure 1.A characterizes…") must stay with
        # its table. Releasing it stranded the sentence alone at the top of a
        # page with three-quarters of that page blank and the table overleaf —
        # worse than the gap the release was meant to close. Only release
        # paragraphs that are NOT introducing this figure.
        prev = tbl._tbl.getprevious()
        released = 0
        while prev is not None and prev.tag == qn('w:p') and released < 2:
            pPr = prev.find(qn('w:pPr'))
            txt = "".join(prev.itertext()).strip()
            is_lead_in = bool(re.match(r'^(Figure|Table)\s+[0-9]', txt))
            if pPr is not None and not is_lead_in:
                for kn in pPr.findall(qn('w:keepNext')):
                    pPr.remove(kn)
                kn = OxmlElement('w:keepNext'); kn.set(qn('w:val'), 'false')
                pPr.append(kn)          # override the style's keepNext
            elif is_lead_in and pPr is not None:
                # glue it to the table instead
                for kn in pPr.findall(qn('w:keepNext')):
                    pPr.remove(kn)
                pPr.append(OxmlElement('w:keepNext'))
            prev = prev.getprevious()
            released += 1

        # 4. a table that breaks should leave at least two rows on each side of
        #    the break rather than one orphaned row.
        if len(rows) >= 4:
            for r in (rows[1], rows[-1]):
                trPr = r._tr.get_or_add_trPr()
                if trPr.find(qn('w:cantSplit')) is None:
                    trPr.append(OxmlElement('w:cantSplit'))

def shrink_oversized_tables():
    """Reclaim vertical space in tall tables so they leave smaller gaps.

    Height model: a row is as tall as its TALLEST cell, and a cell's height is
    driven by how many lines its text wraps to at that column's width. An
    earlier version measured only the longest cell per row and reported 10 tall
    tables; measuring every cell against its own column width found 99. That
    undercount is why the first pass barely helped.

    Two levers, applied only to tables that would overflow a page:
      1. Cell padding: trim paragraph space-before/after inside cells to zero
         and tighten line spacing. Invisible to a reader, and on a 40-row table
         it recovers several lines.
      2. Font: drop a half point, floored at 7.5pt so it stays legible.
    """
    PAGE_LINES = 46
    TEXT_DXA   = 9360          # usable text width in twentieths of a point

    def col_widths(t):
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is None:
            n = len(t.columns) or 1
            return [TEXT_DXA / n] * n
        w = []
        for gc in grid.findall(qn('w:gridCol')):
            try: w.append(float(gc.get(qn('w:w'))))
            except (TypeError, ValueError): w.append(TEXT_DXA / max(1, len(t.columns)))
        return w or [TEXT_DXA]

    def est_lines(t, pt=9.0):
        widths = col_widths(t)
        # ~0.5 * font size per character in twentieths of a point
        char_dxa = pt * 10.0
        total = 0
        for row in t.rows:
            tallest = 1
            for i, cell in enumerate(row.cells):
                w = widths[i] if i < len(widths) else widths[-1]
                per_line = max(8, int(w / char_dxa))
                lines = 0
                for para in cell.paragraphs:
                    lines += max(1, -(-len(para.text) // per_line))
                tallest = max(tallest, lines)
            total += tallest
        return total

    tightened = 0
    for t in doc.tables:
        if est_lines(t) <= PAGE_LINES:
            continue
        tightened += 1
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    pf.space_before = Pt(0); pf.space_after = Pt(0)
                    pf.line_spacing = 1.0
                    for r in para.runs:
                        cur = r.font.size
                        r.font.size = Pt(max(7.5, (cur.pt if cur else 9.0) - 0.5))
    return tightened

def keep_lists_together():
    """Stop bullet lists from splitting across a page with a big void: set
    keepNext on every list item except the last in each run, and keepLines on
    each item. Word then keeps the list on one page (pushing it whole to the
    next page if it doesn't fit) instead of leaving a gap mid-list.
    Skips lists inside callout boxes (already contained)."""
    paras = doc.paragraphs
    n=len(paras)
    i=0
    while i<n:
        if paras[i]._p.find('.//'+qn('w:numPr')) is not None:
            j=i
            while j<n and paras[j]._p.find('.//'+qn('w:numPr')) is not None:
                j+=1
            run=paras[i:j]
            # keepLines on every item (an item won't split mid-itself).
            for p in run:
                pPr=p._p.get_or_add_pPr()
                if pPr.find(qn('w:keepLines')) is None:
                    el=OxmlElement('w:keepLines'); el.set(qn('w:val'),'true'); pPr.append(el)
            # Only keep the WHOLE list together if it's SHORT (<=6 items). Long
            # lists are allowed to flow across a page break naturally (no big void).
            if len(run) <= 6:
                for p in run[:-1]:
                    pPr=p._p.get_or_add_pPr()
                    if pPr.find(qn('w:keepNext')) is None:
                        el=OxmlElement('w:keepNext'); el.set(qn('w:val'),'true'); pPr.append(el)
            else:
                # keep only the FIRST two items glued to the heading above; rest flows
                for p in run[:2]:
                    pPr=p._p.get_or_add_pPr()
                    if pPr.find(qn('w:keepNext')) is None:
                        el=OxmlElement('w:keepNext'); el.set(qn('w:val'),'true'); pPr.append(el)
            i=j
        else:
            i+=1

# ── run pipeline (order matters) ─────────────────────────────────────────────
shrink_sources()          # FIRST: demote 'Sources' headings to Normal BEFORE bookmarking,
                          # so they never get a _Toc bookmark or a TOC entry.
strip_heading_bookmarks()
number_and_color_headings()
style_chapter_abstracts()
style_stat_strips()
style_quotes()
style_depmap()
style_banners()
style_part_divider()
style_callouts()
style_tables()
style_dependency_matrix()
style_captions()
style_inline_figures()
wide_tables_to_landscape()
style_key_concepts()
rule_above_major_sections()
tighten_table_pagination()
shrink_oversized_tables()
keep_lists_together()
chapter_page_breaks()
add_cover()
insert_toc()
add_page_numbers()        # AFTER insert_toc: splits on the TOC's trailing break
finalize_sources()        # LAST: authoritative look for every Sources block

# update fields on open
doc.settings.element.append(_uf := OxmlElement('w:updateFields')); _uf.set(qn('w:val'),'true')

doc.save(OUT)
print(f"Wrote {OUT}")
