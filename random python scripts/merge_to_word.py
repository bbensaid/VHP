import os
from docx import Document
from docx.shared import Inches

# --- CONFIGURATION ---
SOURCE_FOLDER = 'book-images'
OUTPUT_FILENAME = 'continuous_images.docx'
# ---------------------

def create_word_doc():
    # 1. Check folder
    if not os.path.exists(SOURCE_FOLDER):
        print(f"ERROR: Folder '{SOURCE_FOLDER}' not found.")
        return

    # 2. Get files sorted by time
    files = [os.path.join(SOURCE_FOLDER, f) for f in os.listdir(SOURCE_FOLDER) if f.lower().endswith('.png')]
    files.sort(key=os.path.getmtime)

    if not files:
        print("No PNG files found.")
        return

    print(f"Combining {len(files)} images into Word Doc...")

    # 3. Create Document
    doc = Document()
    
    # Set Narrow Margins (0.5 inch) so images are larger and fit more per page
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # 4. Add images one by one
    for file_path in files:
        try:
            # Add the image, forcing it to be 7.5 inches wide (fits perfectly in margins)
            # This auto-scales the height so it doesn't stretch.
            doc.add_picture(file_path, width=Inches(7.5))
            
            # Optional: Ensure no extra paragraph spacing is forcing gaps
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.paragraph_format.space_after = 0
            
        except Exception as e:
            print(f"Could not add {file_path}: {e}")

    # 5. Save
    doc.save(OUTPUT_FILENAME)
    print(f"DONE. Saved as '{OUTPUT_FILENAME}'")
    print("Open this file in Word (or Pages/LibreOffice) to print.")

if __name__ == "__main__":
    create_word_doc()