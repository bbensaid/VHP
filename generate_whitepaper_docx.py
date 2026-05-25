#!/usr/bin/env python3
"""
Generate HTR_WHITE_PAPER.docx from HTR_WHITE_PAPER.md
Produces a properly formatted Word document with real tables, styled headings,
and styled text boxes for diagrams — opens cleanly in Google Docs.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colors ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(15,  40,  80)
SKY    = RGBColor(14,  165, 233)
EMERALD= RGBColor(16,  185, 129)
INDIGO = RGBColor(99,  102, 241)
VIOLET = RGBColor(139, 92,  246)
TEAL   = RGBColor(20,  184, 166)
LIGHT_GRAY = RGBColor(245, 247, 250)
MID_GRAY   = RGBColor(229, 231, 235)
TABLE_HDR  = RGBColor(15,  40,  80)
TABLE_ALT  = RGBColor(240, 244, 250)
DIAGRAM_BG = RGBColor(235, 240, 252)
DIAGRAM_BORDER = RGBColor(99, 102, 241)

# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_to_hex(color):
    """Convert RGBColor to 6-char uppercase hex string."""
    return str(color).upper()

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for k, v in kwargs[edge].items():
                tag.set(qn(f'w:{k}'), str(v))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_run_font(run, size_pt, bold=False, color=None, italic=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def set_para_spacing(para, before=0, after=6, line_rule=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_rule:
        pf.line_spacing_rule = line_rule

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '9CA3AF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_cover_page(doc, title, subtitle, date_str):
    """Full-page cover: navy background, white text."""
    # Title block
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HEALTH TRANSFORMATION REVIEW")
    set_run_font(run, 13, bold=True, color=SKY)
    set_para_spacing(p, 0, 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 32, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, 16, color=RGBColor(60, 80, 120), italic=True)
    set_para_spacing(p, 0, 20)

    add_horizontal_rule(doc)

    meta_items = [
        ("Version", "1.0"),
        ("Date", date_str),
        ("Audience", "Institutional Partners, State Agency Leaders, Policy Professionals"),
        ("Framework", "Six-Pillar Health Transformation Model"),
        ("Primary Case", "Vermont — Acts 167 (2022) and 68 (2025)"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}:  ")
        set_run_font(run, 11, bold=True, color=NAVY)
        run2 = p.add_run(value)
        set_run_font(run2, 11, color=RGBColor(40, 60, 100))
        set_para_spacing(p, 2, 2)

    doc.add_page_break()


def add_toc_page(doc, toc_lines):
    p = doc.add_paragraph()
    run = p.add_run("TABLE OF CONTENTS")
    set_run_font(run, 18, bold=True, color=NAVY)
    set_para_spacing(p, 0, 12)
    add_horizontal_rule(doc)

    for line in toc_lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        indent = 0
        if line.startswith("   ") or line.startswith("\t"):
            indent = 1
            line = line.lstrip()
        run = p.add_run(line)
        if re.match(r'^\d+\.', line) and not indent:
            set_run_font(run, 11, bold=True, color=NAVY)
            p.paragraph_format.left_indent = Inches(0)
        else:
            set_run_font(run, 10, color=RGBColor(60, 80, 120))
            p.paragraph_format.left_indent = Inches(0.3)
        set_para_spacing(p, 1, 1)

    doc.add_page_break()


def add_md_table(doc, md_lines):
    """Parse and render a markdown table as a proper Word table."""
    rows = []
    for line in md_lines:
        line = line.strip()
        if not line or re.match(r'^\|[-| :]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    # Pad rows to max_cols
    rows = [r + [''] * (max_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    col_width = Inches(6.5 / max_cols)

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        row.height = Pt(22) if i == 0 else Pt(18)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(cell_text)
            if i == 0:
                # Header row
                set_run_font(run, 10, bold=True, color=RGBColor(255, 255, 255))
                set_cell_bg(cell, TABLE_HDR)
            elif i % 2 == 0:
                set_run_font(run, 9.5)
                set_cell_bg(cell, TABLE_ALT)
            else:
                set_run_font(run, 9.5)
            # Thin borders
            set_cell_border(cell,
                top={'val':'single','sz':'4','color':'D1D5DB'},
                bottom={'val':'single','sz':'4','color':'D1D5DB'},
                left={'val':'single','sz':'4','color':'D1D5DB'},
                right={'val':'single','sz':'4','color':'D1D5DB'})

    doc.add_paragraph()  # spacing after table


def add_diagram_block(doc, lines):
    """Render ASCII/text diagram as a styled indented block with border."""
    content = '\n'.join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.3)
    p.paragraph_format.right_indent  = Inches(0.3)
    p.paragraph_format.space_before  = Pt(8)
    p.paragraph_format.space_after   = Pt(8)

    # Light blue shaded background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FD')
    pPr.append(shd)

    # Border around paragraph
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'left', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '6366F1')
        pBdr.append(el)
    pPr.append(pBdr)

    run = p.add_run(content)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(15, 40, 80)
    return p


def add_heading1(doc, text):
    """Major section heading (##)."""
    # Page break before major sections (except first)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Colored left border via shading isn't direct, use custom paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '0F2850')
    pBdr.append(left)
    pPr.append(pBdr)

    run = p.add_run(text.upper())
    set_run_font(run, 16, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.15)

    add_horizontal_rule(doc)
    return p


def add_heading2(doc, text):
    """Subsection heading (###)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 13, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_heading3(doc, text):
    """Sub-subsection heading (####)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 11, bold=True, color=RGBColor(40, 80, 140))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p


def parse_inline(para, text):
    """Add a run with inline bold/italic parsed."""
    # Split on ** and * for bold/italic
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = Pt(10.5)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.font.italic = True
            run.font.size = Pt(10.5)
        else:
            run = para.add_run(part)
            run.font.size = Pt(10.5)


def add_body_para(doc, text, bullet=False, numbered=False, num=None, indent=0):
    """Normal body paragraph with optional bullet."""
    p = doc.add_paragraph()
    if bullet:
        p.paragraph_format.left_indent   = Inches(0.3 + indent * 0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        run0 = p.add_run("• ")
        run0.font.color.rgb = NAVY
        run0.font.size = Pt(10.5)
        run0.font.bold = True
    elif numbered and num is not None:
        p.paragraph_format.left_indent   = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run0 = p.add_run(f"{num}. ")
        run0.font.color.rgb = NAVY
        run0.font.size = Pt(10.5)
        run0.font.bold = True
    else:
        p.paragraph_format.left_indent = Inches(indent * 0.25)

    parse_inline(p, text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_callout_box(doc, text, color=INDIGO):
    """Highlighted callout / note box."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4FF')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        pBdr.append(el)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), rgb_to_hex(color))
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = NAVY
    return p


# ── Main parser ───────────────────────────────────────────────────────────────

def process_md(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    in_toc = False
    toc_lines = []
    cover_done = False
    num_counters = {}  # track numbered list state per indent level

    while i < len(lines):
        line = lines[i]

        # ── Title (first H1) → cover page
        if line.startswith('# ') and not cover_done:
            title = line[2:].strip()
            subtitle = ''
            # Peek at next lines for subtitle
            if i+1 < len(lines) and lines[i+1].startswith('## '):
                subtitle = lines[i+1][3:].strip()
                i += 1
            add_cover_page(doc, title, subtitle, "May 2026")
            cover_done = True
            i += 1
            continue

        # ── Front matter metadata lines (** lines at top)
        if re.match(r'^\*\*[^*]+\*\*', line) and not cover_done:
            i += 1
            continue

        # ── TABLE OF CONTENTS block
        if line.strip() == '## TABLE OF CONTENTS':
            in_toc = True
            i += 1
            continue

        if in_toc:
            if line.startswith('## ') or line.startswith('# '):
                # TOC ended
                in_toc = False
                add_toc_page(doc, toc_lines)
                # Don't skip this line — fall through
            elif line.strip() == '---':
                in_toc = False
                add_toc_page(doc, toc_lines)
                i += 1
                continue
            else:
                toc_lines.append(line)
                i += 1
                continue

        # ── Horizontal rule
        if line.strip() == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Page break markers
        if line.strip() == '\\newpage' or line.strip() == '<div style="page-break-after: always;"></div>':
            doc.add_page_break()
            i += 1
            continue

        # ── H2 — Major section headings
        if line.startswith('## '):
            text = line[3:].strip()
            add_heading1(doc, text)
            num_counters = {}
            i += 1
            continue

        # ── H3 — Subsection
        if line.startswith('### '):
            text = line[4:].strip()
            add_heading2(doc, text)
            num_counters = {}
            i += 1
            continue

        # ── H4 — Sub-subsection
        if line.startswith('#### '):
            text = line[5:].strip()
            add_heading3(doc, text)
            i += 1
            continue

        # ── Markdown table
        if line.startswith('|') and i+1 < len(lines) and re.match(r'^\|[-| :]+\|', lines[i+1]):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_md_table(doc, table_lines)
            continue

        # ── Code block / ASCII diagram
        if line.strip().startswith('```'):
            diag_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                diag_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if diag_lines:
                add_diagram_block(doc, diag_lines)
            continue

        # ── Bullet list
        if re.match(r'^[-*] ', line):
            text = line[2:].strip()
            # Check for bold label at start
            add_body_para(doc, text, bullet=True)
            num_counters = {}
            i += 1
            continue

        # ── Indented bullet
        if re.match(r'^  [-*] ', line):
            text = line[4:].strip()
            add_body_para(doc, text, bullet=True, indent=1)
            i += 1
            continue

        # ── Numbered list
        m = re.match(r'^(\d+)\. (.+)', line)
        if m:
            num = int(m.group(1))
            text = m.group(2).strip()
            add_body_para(doc, text, numbered=True, num=num)
            i += 1
            continue

        # ── Bold standalone line (label heading)
        if re.match(r'^\*\*[^*]+\*\*$', line.strip()) or re.match(r'^\*\*[^*]+\*\* —', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            parse_inline(p, line.strip())
            for run in p.runs:
                if run.bold:
                    run.font.size = Pt(11)
                    run.font.color.rgb = NAVY
            i += 1
            continue

        # ── Empty line
        if not line.strip():
            i += 1
            continue

        # ── Normal body paragraph
        if line.strip():
            add_body_para(doc, line.strip())
            i += 1
            continue

        i += 1


def set_document_margins(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1.1)
    section.right_margin = Inches(1.1)
    section.top_margin   = Inches(1.0)
    section.bottom_margin= Inches(1.0)


def set_default_font(doc):
    from docx.oxml.ns import qn
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(20, 20, 40)
    normal.paragraph_format.line_spacing = Pt(15)


def add_header_footer(doc, title_short):
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"HTR White Paper  |  {title_short}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 140, 180)

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_left = fp.add_run("Health Transformation Review  |  healthtransformationreview.com")
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = RGBColor(120, 140, 180)
    fp.add_run("          ")
    # Page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg = fp.add_run()
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)
    run_pg.font.size = Pt(8)
    run_pg.font.color.rgb = RGBColor(120, 140, 180)


def main():
    src = Path('/Users/baba/Vermont-Health-Platform/HTR_WHITE_PAPER.md')
    out = Path('/Users/baba/Vermont-Health-Platform/HTR_WHITE_PAPER.docx')

    md_text = src.read_text(encoding='utf-8')

    doc = Document()
    set_document_margins(doc)
    set_default_font(doc)
    add_header_footer(doc, "Transforming American Healthcare")

    process_md(doc, md_text)

    doc.save(str(out))
    print(f"✓ Saved: {out}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")


if __name__ == '__main__':
    main()
